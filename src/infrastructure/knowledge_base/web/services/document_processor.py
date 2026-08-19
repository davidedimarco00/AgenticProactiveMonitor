from pathlib import Path
import codecs

import fitz
from docx import Document


SUPPORTED_EXTENSIONS = {
    "pdf",
    "docx",
    "txt",
    "md",
}


def extract_text(
        file_path: str,
) -> str:
    path = Path(file_path)

    extension = (
        path.suffix
        .lower()
        .lstrip(".")
    )

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: .{extension}"
        )

    if extension == "pdf":
        text = _extract_pdf(path)

    elif extension == "docx":
        text = _extract_docx(path)

    else:
        text = _extract_plain_text(path)

    text = _normalise_text(text)

    if not text:
        raise ValueError(
            "No readable text was found in the document."
        )

    return text


def chunk_text(
        text: str,
        chunk_size_words: int = 220,
        overlap_words: int = 40,
) -> list[str]:

    if chunk_size_words <= 0:
        raise ValueError(
            "chunk_size_words must be greater than zero."
        )

    if (
            overlap_words < 0
            or overlap_words >= chunk_size_words
    ):
        raise ValueError(
            "overlap_words must be >= 0 "
            "and smaller than chunk_size_words."
        )

    words = text.split()

    if not words:
        return []

    step = (
            chunk_size_words
            - overlap_words
    )

    chunks: list[str] = []

    for start in range(
            0,
            len(words),
            step,
    ):
        end = min(
            start + chunk_size_words,
            len(words),
            )

        chunk = " ".join(
            words[start:end]
        ).strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(words):
            break

    return chunks


def _extract_pdf(
        path: Path,
) -> str:

    parts: list[str] = []

    with fitz.open(path) as pdf:

        for page in pdf:

            page_text = page.get_text(
                "text"
            )

            if page_text:
                parts.append(
                    page_text
                )

    return "\n".join(parts)


def _extract_docx(
        path: Path,
) -> str:

    document = Document(path)

    parts = [
        paragraph.text
        for paragraph
        in document.paragraphs
        if paragraph.text.strip()
    ]

    for table in document.tables:

        for row in table.rows:

            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                parts.append(
                    " | ".join(cells)
                )

    return "\n".join(parts)


def _extract_plain_text(
        path: Path,
) -> str:
    """
    Read TXT/MD files supporting the most common encodings.

    BOM detection is performed first so UTF-16 files are not
    incorrectly decoded as Latin-1.
    """

    raw = path.read_bytes()

    if raw.startswith(
            codecs.BOM_UTF8
    ):
        return raw.decode(
            "utf-8-sig"
        )

    if (
            raw.startswith(codecs.BOM_UTF16_LE)
            or raw.startswith(codecs.BOM_UTF16_BE)
    ):
        return raw.decode(
            "utf-16"
        )

    if (
            raw.startswith(codecs.BOM_UTF32_LE)
            or raw.startswith(codecs.BOM_UTF32_BE)
    ):
        return raw.decode(
            "utf-32"
        )

    try:
        return raw.decode(
            "utf-8"
        )

    except UnicodeDecodeError:
        pass

    try:
        return raw.decode(
            "cp1252"
        )

    except UnicodeDecodeError:
        pass

    return raw.decode(
        "latin-1"
    )


def _normalise_text(
        text: str,
) -> str:

    # Remove NUL characters that should never remain
    # in normal text used for embeddings.
    text = text.replace(
        "\x00",
        ""
    )

    lines = [
        " ".join(
            line.split()
        )
        for line
        in text.splitlines()
    ]

    return "\n".join(
        line
        for line in lines
        if line
    ).strip()